# Тесты python-хелпера doc_to_text.py (xlsx/docx). .xls не генерируем (xlrd только читает) —
# он покрыт ручной проверкой на реальных файлах + проверкой импорта либ в CI (docker-validate).
import os
import subprocess
import sys

SCRIPT = os.path.join(os.path.dirname(__file__), "..", "doc_to_text.py")


def _run(path):
    return subprocess.run(
        [sys.executable, SCRIPT, str(path)], capture_output=True, text=True,
    )


def test_xlsx_extracts_text_and_normalises_numbers(tmp_path):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["Поставщик", "ООО Тест"])
    ws.append(["УНП", 123456789])
    ws.append(["Цена", 1000.0])
    p = tmp_path / "invoice.xlsx"
    wb.save(p)

    r = _run(p)
    assert r.returncode == 0, r.stderr
    assert "ООО Тест" in r.stdout
    assert "123456789" in r.stdout
    # целые числа нормализованы: 1000, а не 1000.0
    assert "1000" in r.stdout and "1000.0" not in r.stdout


def test_docx_extracts_paragraphs_and_tables(tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("Счёт № 7 от 01.06.2026")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Товар"
    t.rows[0].cells[1].text = "Краска MAXIMA"
    p = tmp_path / "invoice.docx"
    d.save(p)

    r = _run(p)
    assert r.returncode == 0, r.stderr
    assert "Счёт № 7" in r.stdout
    assert "Краска MAXIMA" in r.stdout


def test_xlsx_with_doctype_rejected_as_xxe(tmp_path):
    # #57: легитимные xlsx/docx не содержат DTD; файл с <!DOCTYPE/<!ENTITY> в XML-части —
    # XXE-вектор и должен быть отклонён ДО парсера (openpyxl/python-docx на lxml).
    import zipfile
    p = tmp_path / "evil.xlsx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr(
            "xl/worksheets/sheet1.xml",
            '<?xml version="1.0"?>'
            '<!DOCTYPE x [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>'
            "<root>&xxe;</root>",
        )
    r = _run(p)
    assert r.returncode == 1
    assert "DTD/ENTITY" in r.stderr


def test_xlsx_with_utf16_doctype_rejected(tmp_path):
    # #57: гейт не должен обходиться UTF-16-кодировкой XML-части (lxml читает BOM).
    import zipfile
    p = tmp_path / "utf16.xlsx"
    payload = '<?xml version="1.0" encoding="UTF-16"?><!DOCTYPE x [<!ENTITY e SYSTEM "file:///etc/passwd">]><r>&e;</r>'
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr("xl/worksheets/sheet1.xml", b"\xff\xfe" + payload.encode("utf-16-le"))
    r = _run(p)
    assert r.returncode == 1
    assert "DTD/ENTITY" in r.stderr


def test_docx_with_doctype_in_rels_rejected(tmp_path):
    # #57: гейт покрывает и .docx, и .rels-части (не только .xml).
    import zipfile
    p = tmp_path / "evil.docx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr("word/document.xml", "<doc/>")
        z.writestr(
            "_rels/.rels",
            '<!DOCTYPE x [<!ENTITY e SYSTEM "file:///etc/passwd">]><Relationships/>',
        )
    r = _run(p)
    assert r.returncode == 1
    assert "DTD/ENTITY" in r.stderr


def test_unsupported_extension_exits_nonzero(tmp_path):
    p = tmp_path / "note.txt"
    p.write_text("hi")
    r = _run(p)
    assert r.returncode == 1
