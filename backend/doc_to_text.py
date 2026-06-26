#!/usr/bin/env python3
# Извлечение ТЕКСТА из офисных форматов для агента procure-ai.
# Используется backend'ом (backend/extract-text.js) как подпроцесс — это безопаснее,
# чем Node-библиотека SheetJS (неустранённые в npm уязвимости на недоверенных файлах).
#
# Использование:  python3 doc_to_text.py <файл>
# Вывод:          текст документа на stdout
# Коды возврата:  0 — успех, 1 — ошибка/неподдерживаемый формат (детали на stderr)
import os
import sys
import zipfile

# Жёсткий лимит адресного пространства процесса (анти-zip-bomb для xlsx/docx — это ZIP):
# при разбухании получим MemoryError внутри этого процесса, а не OOM-kill всего контейнера.
try:
    import resource
    _cap = int(os.environ.get("DOC_MEM_LIMIT_BYTES", str(512 * 1024 * 1024)))
    resource.setrlimit(resource.RLIMIT_AS, (_cap, _cap))
except Exception:  # noqa: BLE001 — на платформах без RLIMIT просто пропускаем
    pass


def _fmt(value):
    """Аккуратное строковое представление: целые float'ы без «.0» (цена 1000, не 1000.0)."""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


# Маркеры, которых НЕ бывает в легитимных xlsx/docx (это ZIP с XML-частями): DOCTYPE/ENTITY —
# признак XXE / «billion laughs». Блокируем на входе, не полагаясь на безопасные дефолты lxml,
# который используют openpyxl/python-docx и в который нельзя надёжно внедрить безопасный парсер (#57).
_XXE_MARKERS = (b"<!doctype", b"<!entity")

# Сигнатуры контейнеров: Office Open XML (xlsx/docx) — это ZIP (`PK\x03\x04`), а старый .xls —
# OLE2 compound file (`\xd0\xcf\x11\xe0…`). Расширение часто врёт: ERP/1С отдают xlsx под именем
# `.xls` (и наоборот). Если выбрать парсер только по расширению, xlrd падает на zip-«.xls» →
# извлечение возвращает пусто → агент уходит в `unreadable_document` и теряет ВСЕ листы (#334).
_ZIP_MAGIC = b"PK\x03\x04"
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _sniff(path):
    """Определить реальный контейнер по первым байтам: 'zip' | 'ole2' | None."""
    try:
        with open(path, "rb") as f:
            head = f.read(8)
    except OSError:
        return None
    if head.startswith(_ZIP_MAGIC):
        return "zip"
    if head.startswith(_OLE2_MAGIC):
        return "ole2"
    return None


def _reject_xxe_in_zip(path):
    """Просканировать XML-части офисного ZIP на DTD/ENTITY и отклонить файл при их наличии.
    DOCTYPE по правилам XML стоит в начале части, поэтому читаем ограниченный префикс. Значения
    ячеек хранятся XML-экранированными (`&lt;!DOCTYPE`), так что ложных срабатываний на контенте нет."""
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if not name.lower().endswith((".xml", ".rels")):
                continue
            with z.open(name) as f:
                # Read a generous prefix (DOCTYPE must precede the root element by XML spec).
                # Strip NUL bytes BEFORE matching: a UTF-16-encoded part renders "<!DOCTYPE" as
                # "<\x00!\x00D…" (LE) / "\x00<\x00!…" (BE) — removing NULs reconstructs the ASCII
                # markers so a UTF-16 file can't evade the gate (lxml honours the BOM). Valid UTF-8
                # XML has no NULs, so this is a no-op there.
                head = f.read(1024 * 1024).replace(b"\x00", b"").lower()
            if any(m in head for m in _XXE_MARKERS):
                raise ValueError("DTD/ENTITY not allowed in {} (possible XXE)".format(name))


def xlsx_text(path):
    """Читаем .xlsx через openpyxl (data_only — значения, а не формулы; read_only — потоково)."""
    _reject_xxe_in_zip(path)  # #57: отсечь DTD/ENTITY до парсера
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append("# Лист: {}".format(ws.title))
        for row in ws.iter_rows(values_only=True):
            cells = [_fmt(c) for c in row if c is not None and str(c).strip() != ""]
            if cells:
                out.append("\t".join(cells))
    return "\n".join(out)


def xls_text(path):
    """Читаем старый .xls через xlrd; даты → ISO, целые числа → без «.0»."""
    import xlrd
    wb = xlrd.open_workbook(path)
    out = []
    for sh in wb.sheets():
        out.append("# Лист: {}".format(sh.name))
        for r in range(sh.nrows):
            cells = []
            for c in sh.row(r):
                if c.ctype == xlrd.XL_CELL_EMPTY:
                    continue
                if c.ctype == xlrd.XL_CELL_DATE:
                    try:
                        v = str(xlrd.xldate_as_datetime(c.value, wb.datemode).date())
                    except Exception:  # noqa: BLE001 — битая дата → как есть
                        v = _fmt(c.value)
                elif c.ctype == xlrd.XL_CELL_NUMBER:
                    v = _fmt(c.value)
                else:
                    v = str(c.value)
                if v.strip():
                    cells.append(v)
            if cells:
                out.append("\t".join(cells))
    return "\n".join(out)


def docx_text(path):
    """Читаем .docx через python-docx: абзацы + ячейки таблиц (включая вложенные)."""
    _reject_xxe_in_zip(path)  # #57: отсечь DTD/ENTITY до парсера
    import docx

    def cell_text(cell):
        parts = [p.text for p in cell.paragraphs if p.text.strip()]
        for nested in cell.tables:  # вложенные таблицы
            for row in nested.rows:
                rc = [c.text.strip() for c in row.cells if c.text.strip()]
                if rc:
                    parts.append("\t".join(rc))
        return "\n".join(parts)

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [cell_text(c).strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) != 2:
        print("usage: doc_to_text.py <file>", file=sys.stderr)
        return 1
    path = sys.argv[1]
    ext = os.path.splitext(path)[1].lower()
    handlers = {".xlsx": xlsx_text, ".xls": xls_text, ".docx": docx_text}
    handler = handlers.get(ext)
    if handler is None:
        print("unsupported extension: {}".format(ext), file=sys.stderr)
        return 1
    # #334: для таблиц (.xls/.xlsx) выбираем парсер по СИГНАТУРЕ, а не по расширению — оно врёт.
    # ZIP → openpyxl (xlsx, все листы), OLE2 → xlrd (старый xls, все листы). .docx тоже ZIP, но по
    # расширению однозначен — его не переопределяем. Неизвестная сигнатура → доверяем расширению.
    if ext in (".xls", ".xlsx"):
        magic = _sniff(path)
        if magic == "zip":
            handler = xlsx_text
        elif magic == "ole2":
            handler = xls_text
    try:
        sys.stdout.write(handler(path))
        return 0
    except Exception as e:  # noqa: BLE001 — любой сбой парсинга → ошибка извлечения
        print("extract error: {}".format(e), file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
